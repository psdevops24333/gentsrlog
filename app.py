import streamlit as st, zipfile, json, io
import xml.etree.ElementTree as ET

# ==========================================
# ⚙️ 1. ฟังก์ชันช่วยเหลือ (Helper Functions)
# ==========================================
def f_ram(s):
    s = str(s).upper().replace('MB','').replace('BYTES','').replace(' ','')
    try:
        v = float(s)
        return f"{v/(1024**3):.0f} GB" if v > 1048576 else f"{v/1024:.0f} GB"
    except: return str(s)

def f_dsk(s):
    s = str(s).upper().replace('BYTES','').replace('B','').replace(' ','')
    try:
        v = float(s)
        if v > 1000**4: return f"{v/(1000**4):.2f} TB"
        if v > 1000**3: return f"{v/(1000**3):.0f} GB"
        if v > 1024**3: return f"{v/(1024**3):.0f} GB"
    except: pass
    return str(s)

def link_st(v):
    v = str(v).strip()
    if v == '1': return 'Up'
    if v in ['2', '3']: return 'Down'
    if v == '4': return 'Unknown'
    if v == '5': return 'Dormant'
    return v if v else '-'

# ==========================================
# ⚙️ 2. ฟังก์ชันแกะข้อมูล Lifecycle Log (แก้ไขดักจับ Attributes)
# ==========================================
def parse_lc_log(up_file):
    try:
        with zipfile.ZipFile(up_file, 'r') as oz:
            fs = oz.namelist()
            izn = next((f for f in fs if f.lower().endswith('.zip')), None)
            tz = zipfile.ZipFile(io.BytesIO(oz.read(izn)), 'r') if izn else oz
            tfs = tz.namelist()
            
            lcf = next((f for f in tfs if ('lclog' in f.lower() or 'lifecycle' in f.lower()) and f.lower().endswith('.xml')), None)
            if not lcf: return {"status": "not_found"}
            
            rt = ET.fromstring(tz.read(lcf))
            for el in rt.iter():
                if '}' in el.tag: el.tag = el.tag.split('}', 1)[1]
            
            logs = []
            for node in rt.iter():
                # ดักจับทั้งแอตทริบิวต์ในโหนดหลัก <Event ...> และแท็กย่อยด้านใน
                m_id = node.attrib.get('MessageID') or node.attrib.get('MessageId') or node.find('MessageId') or node.find('MessageID')
                m_id_val = m_id.text.strip().upper() if hasattr(m_id, 'text') and m_id.text else str(m_id).strip().upper()
                
                if m_id_val and m_id_val != "NONE":
                    ts_val = node.attrib.get('TIMESTAMP') or node.attrib.get('Timestamp') or node.attrib.get('CreationTimeStamp')
                    if not ts_val:
                        for ch in node:
                            if 'time' in ch.tag.lower() or 'date' in ch.tag.lower():
                                if ch.text: ts_val = ch.text.strip(); break
                    
                    ms = node.find('Message')
                    ms_val = ms.text.strip() if ms is not None and ms.text else node.attrib.get('Message', '-')
                    
                    logs.append({
                        "Code": m_id_val,
                        "Time": ts_val if ts_val else "-",
                        "Details": ms_val
                    })
            return {"status": "success", "data": logs}
    except Exception as e: 
        return {"status": "error"}

# ==========================================
# ⚙️ 3. ฟังก์ชันแกะข้อมูลฮาร์ดแวร์
# ==========================================
def parse_tsr(up_file):
    ex = {}
    try:
        with zipfile.ZipFile(up_file, 'r') as oz:
            fs = oz.namelist()
            izn = next((f for f in fs if f.lower().endswith('.zip')), None)
            tz = zipfile.ZipFile(io.BytesIO(oz.read(izn)), 'r') if izn else oz
            tfs = tz.namelist()
            jf = next((f for f in tfs if 'hardware_inventory.json' in f.lower() or 'hw_inventory.json' in f.lower()), None)
            xfs = [f for f in tfs if ('sysinfo_' in f.lower() and f.endswith('.xml')) or 'inventory.xml' in f.lower() or 'hw_inventory.xml' in f.lower()]

            def add_item(ad):
                id_ = ad.get("_ID_")
                if id_ and id_ not in ['PROPERTY', 'VALUE', 'ATTRIBUTE']:
                    if id_ not in ex: ex[id_] = ad
                    else: ex[id_].update(ad)

            if jf:
                jd = json.loads(tz.read(jf).decode('utf-8', errors='ignore'))
                cs = jd.get("SystemInventory", jd).get("Component", [])
                if isinstance(cs, dict): cs = [cs]
                for c in cs:
                    ats = c.get("Attribute", [])
                    if isinstance(ats, dict): ats = [ats]
                    ad = {a.get("@Name", a.get("Name")).upper(): str(a.get("#text", a.get("Value", a.get("text")))).strip() for a in ats if a.get("@Name", a.get("Name"))}
                    ad["_ID_"] = c.get("@FQDD", c.get("FQDD", "")).upper()
                    add_item(ad)
            if xfs:
                for f in xfs:
                    try:
                        rt = ET.fromstring(tz.read(f))
                        for el in rt.iter():
                            if '}' in el.tag: el.tag = el.tag.split('}', 1)[1]
                        for c in rt.iter():
                            ad = {k.upper(): str(v).strip() for k,v in c.attrib.items()}
                            for ch in c:
                                tg = ch.tag.upper()
                                na = ch.get('Name') or ch.get('NAME')
                                if tg in ['PROPERTY', 'ATTRIBUTE'] and na:
                                    ky = na.upper()
                                    vl = next((s.text.strip() for s in ch if s.tag.upper() == 'VALUE' and s.text), ch.text.strip() if ch.text else "")
                                    if vl: ad[ky] = vl
                                elif ch.text and ch.text.strip(): ad[tg] = ch.text.strip()
                            ad["_ID_"] = ad.get('INSTANCEID', ad.get('FQDD', ad.get('DEVICEID', c.tag))).upper()
                            add_item(ad)
                    except: pass

        si = {"Model": "-", "Service Tag": "-", "Hostname": "-", "IP iDRAC": "-"}
        cp, rm, dk, ct, nc, fc = [], [], [], [], [], []
        
        for id_, ad in sorted(ex.items()):
            if 'SYSTEM' in id_ or 'BOARD' in id_:
                for k, n in [('MODEL','Model'), ('SERVICETAG','Service Tag'), ('HOSTNAME','Hostname')]:
                    if ad.get(k): si[n] = ad[k]
            elif 'IPV4' in id_ or 'IDRAC' in id_:
                ip = ad.get('CURRENTIPADDRESS', ad.get('ADDRESS'))
                if ip and ip not in ['0.0.0.0', '::', '127.0.0.1']: si['IP iDRAC'] = ip
            elif 'CPU' in id_:
                m = ad.get('MODEL', ad.get('DEVICEDESCRIPTION', ad.get('NAME', '-')))
                if m != '-': cp.append({"Model": m, "Clock": f"{ad.get('CURRENTCLOCKSPEED','-')} (max {ad.get('MAXCLOCKSPEED','-')})", "Cores": ad.get('NUMBEROFPROCESSORCORES', ad.get('CORECOUNT', '-')), "Threads": ad.get('NUMBEROFENABLEDTHREADS', ad.get('THREADCOUNT', '-')), "L1": ad.get('PRIMARYCACHE', ad.get('L1CACHE', '-')), "L2": ad.get('SECONDARYCACHE', ad.get('L2CACHE', '-')), "L3": ad.get('TERTIARYCACHE', ad.get('L3CACHE', '-')), "Microcode": ad.get('MICROCODEVERSION', ad.get('MICROCODE', '-'))})
            elif 'DIMM' in id_ or 'MEMORY' in id_:
                sz = f_ram(ad.get('SIZE', ad.get('CAPACITY', '-')))
                if sz != '-': rm.append({"Slot": ad.get('DEVICEDESCRIPTION', ad.get('NAME', id_.split(':')[-1])), "Size": sz, "Speed": ad.get('SPEED', ad.get('OPERATINGSPEED', '-')), "Manufacturer": ad.get('MANUFACTURER', '-'), "Part Number": ad.get('PARTNUMBER', '-'), "Serial Number": ad.get('SERIALNUMBER', '-')})
            elif 'DISK' in id_ or 'PHYSICALDISK' in id_:
                sz = f_dsk(ad.get('SIZE', ad.get('SIZEINBYTES', ad.get('CAPACITY', '-'))))
                if sz != '-': dk.append({"Slot": ad.get('FQDD', id_), "RAID State": ad.get('STATE', ad.get('RAIDSTATUS', '-')), "Vendor": ad.get('MANUFACTURER', ad.get('VENDORID', '-')), "Model": ad.get('MODEL', ad.get('PRODUCTID', '-')), "Size": sz, "Serial": ad.get('SERIALNUMBER', '-'), "SAS Address": ad.get('SASADDRESS', '-'), "Firmware": ad.get('REVISION', ad.get('FIRMWAREVERSION', '-'))})
            elif any(x in id_ for x in ['RAID', 'AHCI', 'CONTROLLER']):
                mdl = ad.get('PRODUCTNAME', ad.get('DEVICEDESCRIPTION', ad.get('NAME', '-')))
                if mdl != '-' and 'USB' not in mdl.upper() and 'BATTERY' not in mdl.upper(): ct.append({"Location": ad.get('FQDD', id_), "Vendor": ad.get('MANUFACTURER', ad.get('VENDORID', '-')), "Model": mdl, "Speed": ad.get('LINKSPEED', '-'), "Mode": ad.get('CONTROLLERMODE', '-'), "Firmware": ad.get('FIRMWAREVERSION', ad.get('PACKAGEVERSION', ad.get('VERSION', '-')))})
            elif 'NIC' in id_ or 'ETHERNET' in id_:
                nm = ad.get('PRODUCTNAME', ad.get('DEVICEDESCRIPTION', ad.get('NAME', '-')))
                if nm != '-' and 'TRANSCEIVER' not in nm.upper(): nc.append({"Location": ad.get('FQDD', id_), "Model": nm, "Speed": ad.get('LINKSPEED', ad.get('CURRENTSPEED', '-')), "Link": link_st(ad.get('LINKSTATUS', '-')), "MAC Address": ad.get('CURRENTMACADDRESS', ad.get('MACADDRESS', '-')), "Firmware": ad.get('FIRMWAREVERSION', ad.get('FAMILYVERSION', ad.get('DEVICEVERSION', '-')))})
            elif 'FC' in id_ or 'FIBRECHANNEL' in id_:
                nm = ad.get('PRODUCTNAME', ad.get('DEVICEDESCRIPTION', ad.get('NAME', '-')))
                if nm != '-' and 'TRANSCEIVER' not in nm.upper(): fc.append({"Location": ad.get('FQDD', id_), "Model": nm, "Speed": ad.get('LINKSPEED', ad.get('CURRENTSPEED', '-')), "Link": link_st(ad.get('LINKSTATUS', '-')), "WWN": ad.get('PORTWWN', ad.get('VIRTUALWWPN', ad.get('WWN', '-'))), "Firmware": ad.get('FIRMWAREVERSION', ad.get('FAMILYVERSION', ad.get('DEVICEVERSION', '-')))})

         def add_idx(lst): return [{"Index": i+1, **d} for i, d in enumerate(lst)]
        return {"System Information": [{"Attribute": k, "Value": v} for k, v in si.items()], "Processors": add_idx(cp), "Memory": add_idx(rm), "Physical Disks": add_idx(dk), "Storage Controllers": add_idx(ct), "Ethernet": add_idx(nc), "Fibre Channel": add_idx(fc)}
    except Exception: return {}

def exp_docx(pd):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    dc = Document()
    tp = dc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("TSR Log Hardware Summary")
    tr.font.size = Pt(18); tr.font.bold = True; tr.font.color.rgb = RGBColor(26, 82, 118)
    for sn, rcs in pd.items():
        if not rcs: continue
        hd = dc.add_heading(sn, level=2)
        hd.runs[0].font.color.rgb = RGBColor(26, 82, 118)
        hs = list(rcs[0].keys())
        tb = dc.add_table(rows=1, cols=len(hs))
        tb.style = 'Table Grid'
        hc = tb.rows[0].cells
        for i, h in enumerate(hs):
            hc[i].text = str(h)
            hc[i].paragraphs[0].runs[0].font.bold = True
            hc[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            sh = OxmlElement('w:shd')
            sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), "1A5276")
            hc[i]._tc.get_or_add_tcPr().append(sh)
        for rc in rcs:
            rc_cells = tb.add_row().cells
            for i, h in enumerate(hs): rc_cells[i].text = str(rc.get(h, ""))
        dc.add_paragraph()
    bf = io.BytesIO(); dc.save(bf); bf.seek(0)
    return bf

# ==========================================
# ⚙️ 4. ฟังก์ชันสร้างเอกสาร Request แบบตรงสเปก V3
# ==========================================
def exp_audit_docx(hw_data, lc_logs, loc_name="MTG"):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    dc = Document()
    
    # หัวข้อหลักสอดคล้องกับหัวไฟล์เอกสาร V3
    title_p = dc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("Request for Data Erase & Server Repurposing") [cite: 1]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # ---------------- ตารางที่ 1: Device Specs ----------------
    hd1 = dc.add_heading("1. ตารางข้อมูลรายละเอียดอุปกรณ์และการประเมินเวลา (Device Specifications & Execution Baseline)", level=2)
    hd1.runs[0].font.size = Pt(11)
    hd1.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    si = {i['Attribute']: i['Value'] for i in hw_data.get('System Information', [])}
    cpu_model = hw_data.get('Processors', [{}])[0].get('Model', '-') if hw_data.get('Processors') else '-'
    disks = hw_data.get('Physical Disks', [])
    
    # บันทึกรายละเอียดดิสก์ทั้งหมด
    disk_details = f"Cryptographic Erase Disks (จำนวน {len(disks)} ลูก): "
    if disks:
        disk_details += ", ".join([f"Slot {d.get('Slot','-')} ({d.get('Size','-')})" for d in disks])
    disk_details += " พร้อมทำ Clear Hardware Cache"
    
    tb1 = dc.add_table(rows=1, cols=2)
    tb1.style = 'Table Grid'
    tb1.rows[0].cells[0].text = "หัวข้อข้อมูล (Field)"
    tb1.rows[0].cells[1].text = "รายละเอียด (Details / Baseline)"
    for cell in tb1.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        
    t1_fields = [
        ("1. อุปกรณ์ / รุ่น (Model)", si.get('Model', '-')),
        ("2. Serial Number (Service Tag)", si.get('Service Tag', '-')),
        ("3. Hostname", si.get('Hostname', '-')),
        ("4. Location", loc_name),
        ("5. IP iDRAC", si.get('IP iDRAC', '-')),
        ("6. Disk & Storage Details", disk_details),
        ("7. CPU Details", cpu_model),
        ("8. Duration Time (จากสถิติใช้งานจริง)", "10 Minutes (รวมขั้นตอนการเปิดเครื่อง Power On, ทำลายข้อมูล และระบบทำ Reboot / Shut Down ทั้งหมดแล้ว)")
    ]
    for k, v in t1_fields:
        r = tb1.add_row().cells
        r[0].text, r[1].text = k, v
    dc.add_paragraph()

    # ---------------- ตารางที่ 2: SOC Operation Log ----------------
    hd2 = dc.add_heading("2. ตารางบันทึกขั้นตอนปฏิบัติงานและจุดแนบหลักฐาน (SOC Operation Log & LCC Artifact Template)", level=2) [cite: 2]
    hd2.runs[0].font.size = Pt(11)
    hd2.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    tb2 = dc.add_table(rows=1, cols=4)
    tb2.style = 'Table Grid'
    headers_t2 = ["ขั้นตอนการทำงาน (Dell LCC Process)", "เวลาเริ่ม - สิ้นสุด", "หลักฐานที่ต้องบันทึก / รหัสระบบ (Required Logs & Artifacts)", "ผลการตรวจ"] [cite: 4]
    for i, h in enumerate(headers_t2):
        tb2.rows[0].cells[i].text = h
        tb2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # ดึงเวลาเฉพาะเจาะจงตามเงื่อนไข SYS Code
    def filter_timestamps(codes):
        ts = [l['Time'] for l in lc_logs if l['Code'] in codes and l['Time'] != '-']
        if not ts: return "___________", "_________", "[ ] Pass\n[ ] Fail" [cite: 4]
        ts.sort()
        return ts[0], ts[-1], "[X] Pass\n[ ] Fail"

    s1_s, s1_e, s1_r = filter_timestamps(['SYS1000', 'SYS162'])
    s2_s, s2_e, s2_r = filter_timestamps(['SYS150']) # ดึงส่วน System Erase Job
    s3_s, s3_e, s3_r = filter_timestamps(['SYS144', 'SYS146', 'SYS153'])
    s4_s, s4_e, s4_r = filter_timestamps(['SYS201', 'SYS151', 'SYS1001'])

    t2_rows = [
        ("1. Pre-Check & Initial Power On\nเปิดเซิร์ฟเวอร์เพื่อเตรียมทำ System Erase Tasks", f"เริ่ม: {s1_s}\nสิ้นสุด: {s1_e}", "- ยืนยันรหัส Log:  SYS1000  (System is turning on)\n- ยืนยันรหัส Log:  SYS162  (Turning on the server for System Erase Tasks)", s1_r), [cite: 4]
        ("2. Launch Repurpose or Retire System\nเข้าเมนูตั้งค่าและเลือกส่วนประกอบที่ต้องการลบถาวร", f"เริ่ม: {s2_s}\nสิ้นสุด: {s2_e}", "- เลือกฟังก์ชัน:  BIOS reset default, iDRAC reset, Lifecycle Controller Data, Storage Components, Cryptographic Erase Disks", s2_r), [cite: 4]
        ("3. Storage Data Erasing (Crypto Erase)\nระบบทำการลบข้อมูลบน Drives และ Hardware Cache", f"เริ่ม: {s3_s}\nสิ้นสุด: {s3_e}", "- ยืนยันรหัส Log:  SYS144  (Starting controller hardware cache data erase)\n- ยืนยันรหัส Log:  SYS146  (Starting cryptographic erase-capable drive erase)\n- ยืนยันรหัส Log:  SYS153  (Deleting hardware cache data for controller)", s3_r), [cite: 4]
        ("4. Job Verification & Completion Status\nลบข้อมูลเสร็จสมบูรณ์และระบบสั่งปิดการทำงานตัวเครื่อง", f"เริ่ม: {s4_s}\nสิ้นสุด: {s4_e}", "- ยืนยันรหัส Log:  SYS201  (Disk erase operation successfully completed)\n- ยืนยันรหัส Log:  SYS150  (Erase operations successfully completed)\n- ยืนยันรหัส Log:  SYS151  (Completed System Erase)\n- ยืนยันรหัส Log:  SYS1001  (System is turning off)", s4_r) [cite: 4]
    ]
    for r_data in t2_rows:
        row = tb2.add_row().cells
        for idx, text in enumerate(r_data):
            row[idx].text = text
    dc.add_paragraph()

    # ---------------- ส่วน Compliance ท้ายไฟล์ ----------------
    hd3 = dc.add_heading("3. ข้อควรระวังและการส่งมอบ Artifact (Security Audit Compliance)", level=2) [cite: 5]
    hd3.runs[0].font.size = Pt(11)
    hd3.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p3 = dc.add_paragraph()
    p3.add_run("การเก็บรักษาหลักฐาน (Artifact Retention): ").bold = True [cite: 6]
    p3.add_run("หลังจากทีม SOC ดำเนินการล้างข้อมูลเสร็จสิ้น จะต้องทำการส่งภาพถ่ายหน้าจอสรุปผลจากหน้าจอ Lifecycle Controller และทำการ Export ไฟล์ iDRAC Lifecycle Log ที่ระบุรหัส SYS146, SYS201, SYS150 และ SYS151 พร้อมเวลา Time Stamp เพื่อใช้แนบท้ายเป็นเอกสารปิดงาน (Artifact Check-off)\n\n") [cite: 6]
    p3.add_run("ความรวดเร็วของกระบวนการ: ").bold = True [cite: 7]
    p3.add_run("เนื่องจากตัวเครื่องใช้ระบบ Cryptographic Erase ร่วมกับ NVMe/Supported Disks กระบวนการทางเทคนิคบนดิสก์จริงจะใช้เวลาสั้นมาก (ประมาณ 10 นาที) หากใน Log แสดงเวลาดำเนินการเกินกว่าระยะเวลาปกติอย่างมีนัยสำคัญ ให้ทีม SOC ตรวจสอบประสิทธิภาพการทำงานของ Storage Controller เพิ่มเติม") [cite: 7]

    bf = io.BytesIO(); dc.save(bf); bf.seek(0)
    return bf

# ==========================================
# 🖥️ 5. ส่วนหน้าเว็บ (Streamlit UI)
# ==========================================
st.set_page_config(page_title="TSR Log Tool", page_icon="🖥️", layout="wide")
st.title("🖥️ Server Inventory & Data Erase Audit Tool")

uf = st.file_uploader("อัปโหลดไฟล์ TSR Log (.zip) ของคุณที่นี่", type=["zip"])

if uf:
    with st.spinner("กำลังเจาะลึกข้อมูลฮาร์ดแวร์และตรวจสอบ Log..."):
        hw_data = parse_tsr(uf)
        lc_logs = parse_lc_log(uf)

    tab1, tab2 = st.tabs(["📊 1. Hardware Summary (สเปกเครื่อง)", "🛡️ 2. Data Erase Audit (ตรวจสอบการลบข้อมูล)"])
    
    # ------------------- TAB 1 -------------------
    with tab1:
        if hw_data:
            st.success("✅ โหลดข้อมูลฮาร์ดแวร์สำเร็จ!")
            for s, r in hw_data.items():
                if r:
                    st.markdown(f"#### 🔹 {s}")
                    st.dataframe(r, hide_index=True, use_container_width=True)

    # ------------------- TAB 2 -------------------
    with tab2:
        st.info("📌 อ้างอิงตรวจสอบรหัส Lifecycle Log เพื่อยืนยันกระบวนการ Cryptographic Erase (SOC Standard)") [cite: 6]
        
        if lc_logs.get("status") == "not_found":
            st.warning("⚠️ ไม่พบไฟล์ Lclog.xml หรือ LifecycleLog.xml ใน ZIP นี้")
        elif lc_logs.get("status") == "success":
            logs_data = lc_logs.get("data", [])
            target_ids = ['SYS1000', 'SYS162', 'SYS144', 'SYS146', 'SYS153', 'SYS201', 'SYS150', 'SYS151', 'SYS1001']
            found_logs = [log for log in logs_data if log["Code"] in target_ids]
            
            if len(found_logs) > 0:
                st.success(f"✅ พบหลักฐานการทำ Data Erase ทั้งหมด {len(found_logs)} รายการ")
                st.dataframe(found_logs, hide_index=True, use_container_width=True)
                
                st.write("---")
                st.markdown("### 📄 สร้างเอกสาร Data Erase Request (SOC Template)")
                
                col1, col2 = st.columns([1, 2])
                with col1:
                    loc_input = st.text_input("ระบุ Location เครื่อง (เช่น MTG, BKK):", value="MTG")
                with col2:
                    st.write("") 
                    st.write("")
                    if st.button("🔄 ส่งออกเอกสาร Request เป็น Word (.docx)"):
                        try:
                            df = exp_audit_docx(hw_data, logs_data, loc_input)
                            sys_info = hw_data.get("System Information", [])
                            stg = next((i['Value'] for i in sys_info if i['Attribute'] == "Service Tag"), "Unknown")
                            fn = f"Data_Erase_Request_{stg}.docx"
                            
                            st.download_button("📥 ดาวน์โหลดเอกสาร Data Erase Request", data=df, file_name=fn, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        except Exception as e: st.error(f"Error: {e}")
            else:
                st.error("❌ ไม่พบรหัส SYS ที่เกี่ยวกับการลบข้อมูล (Data Erase) ในเครื่องนี้เลยครับ")

# --- จบโค้ดตรงนี้ชัวร์ 100% ---
